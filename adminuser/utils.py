# employees/utils.py (create a new file)
import random
import string
from django.core.mail import send_mail
from django.template.loader import render_to_string
from django.core.mail import EmailMultiAlternatives
from recruit_management.settings import EMAIL_HOST_USER,DEFAULT_FROM_EMAIL

def generate_random_password():
    chars = string.ascii_lowercase + string.digits + string.punctuation
    return ''.join(random.choice(chars) for i in range(7))

def send_activation_email(employee, site_url, password):
    emailOfSender = DEFAULT_FROM_EMAIL
    subject = 'Login to Your Account on RecruitSmart'
    message = render_to_string('adminuser/activation_email.html', {
        'employee': employee,
        'site_url': site_url,
        'password': password
    })
    emailMessage = EmailMultiAlternatives(subject=subject, body='text_content', from_email=emailOfSender,
                                          to=[employee.user.email, ], reply_to=['info@jmsadvisory.in'])
    emailMessage.attach_alternative(message, "text/html")
    emailMessage.send(fail_silently=False)

    # send_mail(subject, message, '', [employee.email])




# def extract_resume_text(file):
#     text = ""

#     if file.name.endswith(".pdf"):
#         import pdfplumber
#         with pdfplumber.open(file) as pdf:
#             for page in pdf.pages:
#                 text += page.extract_text() or ""

#     elif file.name.endswith(".docx"):
#         import docx
#         doc = docx.Document(file)
#         for para in doc.paragraphs:
#             text += para.text + "\n"

#     return text

def extract_resume_text(file):
    text = ""

    try:
        # -------- PDF --------
        if file.name.lower().endswith(".pdf"):
            import pdfplumber

            try:
                with pdfplumber.open(file) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
            except Exception:
                # fallback parser
                import fitz  # PyMuPDF
                file.seek(0)
                pdf = fitz.open(stream=file.read(), filetype="pdf")
                for page in pdf:
                    text += page.get_text()

        # -------- DOCX --------
        elif file.name.lower().endswith(".docx"):
            import docx
            doc = docx.Document(file)
            for para in doc.paragraphs:
                text += para.text + "\n"

    except Exception as e:
        print("Resume parsing error:", file.name, e)

    return text
